"""Create documents — PDF (fpdf2), Word (.docx), or Markdown — from text content.

Saves to Desktop by default. Returns the saved path.
"""
from __future__ import annotations
import datetime as _dt
from pathlib import Path


def _desktop() -> Path:
    home = Path.home()
    for c in (home / "OneDrive" / "Desktop", home / "Desktop", home):
        if c.exists():
            return c
    return home


def _resolve(path: str, default_ext: str, default_stem: str) -> Path:
    if path:
        p = Path(path).expanduser()
        if p.suffix.lower() != default_ext:
            p = p.with_suffix(default_ext)
        return p
    name = f"{default_stem}_{_dt.datetime.now():%Y%m%d_%H%M%S}{default_ext}"
    return _desktop() / name


def create_markdown(content: str, title: str = "", path: str = "") -> str:
    p = _resolve(path, ".md", "jarvis_note")
    try:
        p.parent.mkdir(parents=True, exist_ok=True)
        body = (f"# {title}\n\n" if title else "") + (content or "")
        p.write_text(body, encoding="utf-8")
        return f"Markdown saved to {p}"
    except Exception as e:
        return f"Markdown creation failed: {e}"


def create_pdf(content: str, title: str = "", path: str = "") -> str:
    try:
        from fpdf import FPDF
    except ImportError:
        return "PDF needs fpdf2 — run: pip install fpdf2"
    p = _resolve(path, ".pdf", "jarvis_report")
    try:
        pdf = FPDF()
        pdf.set_margins(15, 15, 15)
        pdf.set_auto_page_break(auto=True, margin=15)
        pdf.add_page()
        width = pdf.epw   # effective page width (page minus L/R margins)
        if title:
            pdf.set_font("Helvetica", "B", 16)
            pdf.set_x(pdf.l_margin)
            pdf.multi_cell(width, 10, _latin(title))
            pdf.ln(2)
        pdf.set_font("Helvetica", size=12)
        for line in (content or "").split("\n"):
            pdf.set_x(pdf.l_margin)
            pdf.multi_cell(width, 7, _latin(line) if line.strip() else " ")
        p.parent.mkdir(parents=True, exist_ok=True)
        pdf.output(str(p))
        return f"PDF saved to {p}"
    except Exception as e:
        return f"PDF creation failed: {e}"


def _latin(s: str) -> str:
    """fpdf core fonts are Latin-1 only; replace unsupported chars so it never crashes."""
    return (s or "").encode("latin-1", "replace").decode("latin-1")


def create_docx(content: str, title: str = "", path: str = "") -> str:
    try:
        from docx import Document
    except ImportError:
        return "Word docs need python-docx — run: pip install python-docx"
    p = _resolve(path, ".docx", "jarvis_document")
    try:
        doc = Document()
        if title:
            doc.add_heading(title, level=1)
        for para in (content or "").split("\n"):
            doc.add_paragraph(para)
        p.parent.mkdir(parents=True, exist_ok=True)
        doc.save(str(p))
        return f"Word document saved to {p}"
    except Exception as e:
        return f"Word doc creation failed: {e}"
